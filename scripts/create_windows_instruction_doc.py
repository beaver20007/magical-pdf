from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Magical PDF - Windows setup and Codex handoff.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.08):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F7F2")
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.0)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(28, 43, 38)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, after=4)
    p.add_run(text)


doc = Document()

section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(30, 35, 34)

for style_name, size, color in [
    ("Title", 24, RGBColor(20, 30, 28)),
    ("Heading 1", 16, RGBColor(20, 30, 28)),
    ("Heading 2", 13, RGBColor(40, 66, 57)),
]:
    style = styles[style_name]
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color

header = section.header.paragraphs[0]
header.text = "Magical PDF"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header.runs[0].font.name = "Arial"
header.runs[0].font.size = Pt(9)
header.runs[0].font.color.rgb = RGBColor(100, 112, 108)

footer = section.footer.paragraphs[0]
footer.text = "Инструкция для продолжения работы на Windows"
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.name = "Arial"
footer.runs[0].font.size = Pt(9)
footer.runs[0].font.color.rgb = RGBColor(100, 112, 108)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.add_run("Magical PDF: инструкция для Windows и Codex")

subtitle = doc.add_paragraph()
set_paragraph_spacing(subtitle, after=14)
run = subtitle.add_run(
    "Документ для переноса проекта на ПК: как скачать репозиторий, "
    "подготовить окружение, собрать Windows-версию и продолжить работу в Codex."
)
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(86, 99, 95)

doc.add_heading("1. Что уже готово", level=1)
for item in [
    "Проект перенесен в GitHub-репозиторий: https://github.com/beaver20007/magical-pdf",
    "Публичная веб-версия доступна по адресу: https://beaver20007.github.io/magical-pdf/",
    "Основной функционал: PDF превращается в JPEG-страницы, затем собирается обратно в PDF как скан без выделяемого текстового слоя.",
    "Также есть экспорт JPEG-страниц в ZIP.",
    "Версия macOS уже проверена по функционалу.",
]:
    add_bullet(doc, item)

doc.add_heading("2. Что установить на Windows", level=1)
for item in [
    "Git for Windows.",
    "Node.js LTS.",
    "Rust через rustup.",
    "Visual Studio Build Tools 2022.",
    "В Visual Studio Build Tools выбрать компонент Desktop development with C++.",
    "Проверить, что установлены MSVC v143 и Windows SDK.",
    "WebView2 Evergreen Runtime, если его нет в системе.",
]:
    add_bullet(doc, item)

doc.add_heading("3. Как скачать проект на ПК", level=1)
paragraph = doc.add_paragraph(
    "Открыть PowerShell и выполнить команды:"
)
set_paragraph_spacing(paragraph)
add_code_block(
    doc,
    [
        "cd $env:USERPROFILE\\Documents",
        "git clone https://github.com/beaver20007/magical-pdf.git",
        "cd magical-pdf",
    ],
)

doc.add_heading("4. Как запустить веб-версию локально", level=1)
add_code_block(doc, ["npm ci", "npm run dev:web"])
p = doc.add_paragraph(
    "После запуска открыть в браузере адрес:"
)
set_paragraph_spacing(p)
add_code_block(doc, ["http://127.0.0.1:5173/"])

doc.add_heading("5. Как собрать Windows .exe", level=1)
add_code_block(doc, ["npm run build:windows"])
p = doc.add_paragraph(
    "Готовый инсталлятор должен появиться в папке:"
)
set_paragraph_spacing(p)
add_code_block(doc, ["src-tauri\\target\\release\\bundle\\nsis\\"])

doc.add_heading("6. Как продолжить работу в Codex на ПК", level=1)
doc.add_paragraph(
    "В Codex на Windows открыть папку проекта:"
)
add_code_block(doc, ["%USERPROFILE%\\Documents\\magical-pdf"])
doc.add_paragraph(
    "Затем написать Codex следующую команду:"
)
add_code_block(
    doc,
    [
        "Продолжи работу над Magical PDF. Прочитай CODEX_HANDOFF.md и WINDOWS_SETUP.md,",
        "затем помоги собрать Windows .exe.",
    ],
)

doc.add_heading("7. Важные файлы проекта", level=1)
for item in [
    "index.html - основная структура интерфейса.",
    "styles.css - дизайн приложения.",
    "app.js - логика загрузки, преобразования и скачивания файлов.",
    "src-tauri/src/main.rs - нативное сохранение файлов и открытие Finder/Explorer.",
    "src-tauri/tauri.conf.json - настройки приложения, названия, окна и иконок.",
    "CODEX_HANDOFF.md - краткий контекст проекта для Codex.",
    "WINDOWS_SETUP.md - техническая инструкция по Windows-сборке.",
]:
    add_bullet(doc, item)

doc.add_heading("8. Что важно проверить в Windows-версии", level=1)
for item in [
    "Загружается ли PDF через кнопку выбора файла и перетаскивание.",
    "Создается ли новый PDF как скан, без выделяемого текста.",
    "Создается ли ZIP с JPEG-страницами.",
    "Открывается ли Explorer после сохранения файла.",
    "Корректно ли отображается название Magical PDF и иконка приложения.",
]:
    add_bullet(doc, item)

doc.add_paragraph()
note = doc.add_paragraph()
set_paragraph_spacing(note, before=10, after=0)
run = note.add_run(
    "Примечание: сам чат Codex напрямую перенести нельзя. "
    "Для продолжения работы достаточно открыть репозиторий на ПК и дать Codex указание прочитать CODEX_HANDOFF.md."
)
run.bold = True
run.font.color.rgb = RGBColor(40, 66, 57)

doc.save(OUT)
print(OUT)
