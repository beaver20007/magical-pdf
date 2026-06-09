"""DOCX layout layer — masked scan + DrawingML anchored editable text (Word 2010+)."""

from __future__ import annotations

import html
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.shared import Pt

from src.pipeline.cleanup import clean_ocr_text
from src.pipeline.font_estimate import DEFAULT_FONT_NAME, estimate_font_size_pt
from src.pipeline.ir import DocumentIR, ImageBlock, Page, TableBlock, TextBlock
from src.pipeline.text_mask import mask_page_background

_PT_TO_EMU = 12700
_SHAPE_ID = 1000


def _next_id() -> int:
    global _SHAPE_ID
    _SHAPE_ID += 1
    return _SHAPE_ID


def _emu(pt: float) -> int:
    return int(round(pt * _PT_TO_EMU))


def _set_section(section, page: Page) -> None:
    section.page_width = Pt(page.width_pt)
    section.page_height = Pt(page.height_pt)
    section.top_margin = Pt(0)
    section.bottom_margin = Pt(0)
    section.left_margin = Pt(0)
    section.right_margin = Pt(0)


def _bbox_pt(bbox, page: Page) -> tuple[float, float, float, float]:
    left = bbox.x * page.width_pt
    top = bbox.y * page.height_pt
    width = max(14.0, bbox.w * page.width_pt)
    height = max(9.0, bbox.h * page.height_pt)
    return left, top, width, height


def _vml_fallback(
    text: str,
    left_pt: float,
    top_pt: float,
    width_pt: float,
    height_pt: float,
    font_size_pt: float,
    font_name: str,
) -> str:
    safe = html.escape(text, quote=False)
    sz = max(16, min(48, int(round(font_size_pt * 2))))
    return f"""
    <mc:Fallback xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">
      <w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
          xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office">
        <v:shape type="#_x0000_t202" stroked="f" filled="t" fillcolor="white" o:allowincell="f"
          style="position:absolute;margin-left:{left_pt:.2f}pt;margin-top:{top_pt:.2f}pt;
          width:{width_pt:.2f}pt;height:{height_pt:.2f}pt;z-index:251658240;
          mso-position-horizontal-relative:page;mso-position-vertical-relative:page">
          <v:textbox inset="0,0,0,0">
            <w:txbxContent>
              <w:p>
                <w:r>
                  <w:rPr>
                    <w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>
                    <w:sz w:val="{sz}"/>
                  </w:rPr>
                  <w:t xml:space="preserve">{safe}</w:t>
                </w:r>
              </w:p>
            </w:txbxContent>
          </v:textbox>
        </v:shape>
      </w:pict>
    </mc:Fallback>"""


def _anchor_textbox_xml(
    text: str,
    left_pt: float,
    top_pt: float,
    width_pt: float,
    height_pt: float,
    font_size_pt: float,
    font_name: str,
    shape_id: int,
) -> str:
    safe = html.escape(text, quote=False)
    sz = max(16, min(48, int(round(font_size_pt * 2))))
    left_e, top_e = _emu(left_pt), _emu(top_pt)
    w_e, h_e = _emu(width_pt), _emu(height_pt)
    vml = _vml_fallback(text, left_pt, top_pt, width_pt, height_pt, font_size_pt, font_name)
    return f"""
    <mc:AlternateContent xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">
      <mc:Choice Requires="wps">
        <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
            xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
            xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
          <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"
              relativeHeight="251658240" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="page"><wp:posOffset>{left_e}</wp:posOffset></wp:positionH>
            <wp:positionV relativeFrom="page"><wp:posOffset>{top_e}</wp:posOffset></wp:positionV>
            <wp:extent cx="{w_e}" cy="{h_e}"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:wrapNone/>
            <wp:docPr id="{shape_id}" name="TextBox {shape_id}"/>
            <wp:cNvGraphicFramePr>
              <a:graphicFrameLocks noChangeAspect="1"/>
            </wp:cNvGraphicFramePr>
            <a:graphic>
              <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                <wps:wsp>
                  <wps:cNvSpPr txBox="1"/>
                  <wps:spPr>
                    <a:xfrm><a:off x="0" y="0"/><a:ext cx="{w_e}" cy="{h_e}"/></a:xfrm>
                    <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                    <a:solidFill><a:srgbClr val="FFFFFF"/></a:solidFill>
                    <a:ln w="0"><a:noFill/></a:ln>
                  </wps:spPr>
                  <wps:txbx>
                    <w:txbxContent>
                      <w:p>
                        <w:pPr><w:spacing w:before="0" w:after="0"/></w:pPr>
                        <w:r>
                          <w:rPr>
                            <w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>
                            <w:sz w:val="{sz}"/>
                          </w:rPr>
                          <w:t xml:space="preserve">{safe}</w:t>
                        </w:r>
                      </w:p>
                    </w:txbxContent>
                  </wps:txbx>
                  <wps:bodyPr wrap="square" lIns="0" tIns="0" rIns="0" bIns="0"/>
                </wps:wsp>
              </a:graphicData>
            </a:graphic>
          </wp:anchor>
        </w:drawing>
      </mc:Choice>
      {vml}
    </mc:AlternateContent>"""


def _page_text_blocks(ir: DocumentIR, page_index: int) -> list[TextBlock | TableBlock]:
    blocks: list[TextBlock | TableBlock] = []
    for b in ir.blocks:
        if b.type == "page_break":
            continue
        if isinstance(b, ImageBlock):
            continue
        if getattr(b, "page_index", -1) == page_index:
            if isinstance(b, (TextBlock, TableBlock)):
                blocks.append(b)
    blocks.sort(key=lambda x: (x.bbox.y, x.bbox.x))
    return blocks


def emit_docx_layout(
    ir: DocumentIR,
    output_path: Path,
    *,
    page_backgrounds: list[Path],
    mask_text: bool = True,
    mask_dir: Path | None = None,
) -> Path:
    """
    One section per PDF page: masked scan + anchored editable text boxes.
    Stamps/signatures stay in the background raster (not gap-fill images).
    """
    global _SHAPE_ID
    _SHAPE_ID = 1000

    doc = Document()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    work_mask = mask_dir or output_path.parent / "masked_assets"
    work_mask.mkdir(parents=True, exist_ok=True)

    pages = ir.pages or [Page(index=i) for i in range(ir.source.page_count)]
    page_count = ir.source.page_count

    for pi in range(page_count):
        page = pages[pi] if pi < len(pages) else Page(index=pi)
        if pi > 0:
            doc.add_section()
        _set_section(doc.sections[-1], page)

        bg = page_backgrounds[pi] if pi < len(page_backgrounds) else None
        page_texts = [b for b in _page_text_blocks(ir, pi) if isinstance(b, TextBlock)]

        if bg and bg.exists() and mask_text and page_texts:
            masked = work_mask / f"lukoil_page_{pi:03d}.png"
            bg_use = mask_page_background(bg, page_texts, masked)
        else:
            bg_use = bg

        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(0.1)

        run = para.add_run()
        if bg_use and Path(bg_use).exists():
            run.add_picture(str(bg_use), width=Pt(page.width_pt), height=Pt(page.height_pt))

        for block in _page_text_blocks(ir, pi):
            if isinstance(block, TextBlock):
                text = clean_ocr_text(block.text)
                if not text.strip():
                    continue
                left, top, width, height = _bbox_pt(block.bbox, page)
                font_pt = block.font_size_pt or estimate_font_size_pt(block.bbox, page, text)
                font_name = block.font_name or DEFAULT_FONT_NAME
                run._r.append(
                    parse_xml(
                        _anchor_textbox_xml(
                            text, left, top, width, height, font_pt, font_name, _next_id()
                        )
                    )
                )
            elif isinstance(block, TableBlock):
                rows = block.rows
                if not rows:
                    continue
                cols = max(len(r) for r in rows)
                table_text = "\n".join(
                    " | ".join(row[ci] if ci < len(row) else "" for ci in range(cols))
                    for row in rows
                )
                table_text = clean_ocr_text(table_text)
                left, top, width, height = _bbox_pt(block.bbox, page)
                font_pt = estimate_font_size_pt(block.bbox, page, table_text) * 0.85
                run._r.append(
                    parse_xml(
                        _anchor_textbox_xml(
                            table_text,
                            left,
                            top,
                            width,
                            height,
                            font_pt,
                            DEFAULT_FONT_NAME,
                            _next_id(),
                        )
                    )
                )

    doc.save(output_path)
    return output_path
