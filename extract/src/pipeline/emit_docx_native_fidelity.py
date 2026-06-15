"""Premium native PDF → DOCX: masked background + positioned text with PDF fonts."""

from __future__ import annotations

import html
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from src.pipeline.ir import DocumentIR, ImageBlock, Page, TableBlock, TextBlock, TextRun
from src.pipeline.text_mask import mask_page_background

_PT_TO_EMU = 12700
_SHAPE_ID = 2000
_DEFAULT_FONT = "Garamond"
_FALLBACK_FONT = "Times New Roman"
_CYRILLIC_SAFE_FONTS = {"Garamond", "Garamond-Bold"}


def _next_id() -> int:
    global _SHAPE_ID
    _SHAPE_ID += 1
    return _SHAPE_ID


def _emu(pt: float) -> int:
    return int(round(pt * _PT_TO_EMU))


def _set_section(section, page: Page) -> None:
    section.page_width = Pt(page.width_pt)
    section.page_height = Pt(page.height_pt)
    section.top_margin = Pt(0)
    section.bottom_margin = Pt(0)
    section.left_margin = Pt(0)
    section.right_margin = Pt(0)


def _bbox_pt(bbox, page: Page) -> tuple[float, float, float, float]:
    left = bbox.x * page.width_pt
    top = bbox.y * page.height_pt
    width = max(8.0, bbox.w * page.width_pt)
    height = max(7.0, bbox.h * page.height_pt)
    return left, top, width, height


def _word_font(name: str | None, bold: bool) -> str:
    if not name:
        return _FALLBACK_FONT
    base = name.replace("-Bold", "").replace(",Bold", "").strip()
    if not base:
        return _FALLBACK_FONT
    if base in _CYRILLIC_SAFE_FONTS:
        return _FALLBACK_FONT
    return base


def _runs_for_block(block: TextBlock) -> list[TextRun]:
    if block.runs:
        return [r for r in block.runs if r.text]
    if block.text:
        return [
            TextRun(
                text=block.text,
                font_name=block.font_name or _DEFAULT_FONT,
                font_size_pt=block.font_size_pt or 12.0,
                bold="Bold" in (block.font_name or ""),
            )
        ]
    return []


def _runs_xml(runs: list[TextRun]) -> str:
    parts: list[str] = []
    for run in runs:
        if not run.text:
            continue
        safe = html.escape(run.text, quote=False)
        font = _word_font(run.font_name, run.bold)
        size = run.font_size_pt or 12.0
        sz = max(16, min(72, int(round(size * 2))))
        bold_xml = "<w:b/>" if run.bold or "Bold" in (run.font_name or "") else ""
        parts.append(
            f"""<w:r>
              <w:rPr>
                <w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:cs="{font}"/>
                <w:sz w:val="{sz}"/>
                {bold_xml}
              </w:rPr>
              <w:t xml:space="preserve">{safe}</w:t>
            </w:r>"""
        )
    return "".join(parts) if parts else "<w:r><w:t></w:t></w:r>"


_NS_W   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS_WP  = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_NS_A   = "http://schemas.openxmlformats.org/drawingml/2006/main"
_NS_PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"
_NS_R   = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def _anchor_image_xml(
    rel_id: str,
    width_pt: float,
    height_pt: float,
    shape_id: int,
) -> str:
    """Full-page background image anchored at (0,0) behind all content."""
    w_e, h_e = _emu(width_pt), _emu(height_pt)
    return (
        f'<w:drawing'
        f' xmlns:w="{_NS_W}"'
        f' xmlns:wp="{_NS_WP}"'
        f' xmlns:a="{_NS_A}"'
        f' xmlns:pic="{_NS_PIC}"'
        f' xmlns:r="{_NS_R}">'
        f'<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"'
        f' relativeHeight="1" behindDoc="1" locked="1" layoutInCell="1" allowOverlap="0">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{w_e}" cy="{h_e}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:wrapNone/>'
        f'<wp:docPr id="{shape_id}" name="Background{shape_id}"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic>'
        f'<a:graphicData uri="{_NS_PIC}">'
        f'<pic:pic>'
        f'<pic:nvPicPr>'
        f'<pic:cNvPr id="{shape_id}" name="Background"/>'
        f'<pic:cNvPicPr><a:picLocks noChangeAspect="1"/></pic:cNvPicPr>'
        f'</pic:nvPicPr>'
        f'<pic:blipFill>'
        f'<a:blip r:embed="{rel_id}"/>'
        f'<a:stretch><a:fillRect/></a:stretch>'
        f'</pic:blipFill>'
        f'<pic:spPr>'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{w_e}" cy="{h_e}"/></a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'</pic:spPr>'
        f'</pic:pic>'
        f'</a:graphicData>'
        f'</a:graphic>'
        f'</wp:anchor>'
        f'</w:drawing>'
    )


def _grid_cols(col_count: int, col_w_emu: int) -> str:
    return "".join(f'<w:gridCol w:w="{col_w_emu}"/>' for _ in range(col_count))


def _page_table_blocks(ir: DocumentIR, page_index: int) -> list[TableBlock]:
    blocks: list[TableBlock] = []
    for b in ir.blocks:
        if b.type != "table":
            continue
        if getattr(b, "page_index", -1) == page_index:
            blocks.append(b)
    blocks.sort(key=lambda x: (x.bbox.y, x.bbox.x))
    return blocks


def _page_text_blocks(ir: DocumentIR, page_index: int) -> list[TextBlock]:
    blocks: list[TextBlock] = []
    for b in ir.blocks:
        if b.type != "text":
            continue
        if getattr(b, "page_index", -1) == page_index:
            blocks.append(b)
    blocks.sort(key=lambda x: (x.bbox.y, x.bbox.x))
    return blocks


def _row_to_spans(row: list[str | None]) -> list[tuple[str, int]]:
    """Convert a row with None colspan-markers into (text, colspan) pairs."""
    cells: list[tuple[str, int]] = []
    i = 0
    while i < len(row):
        val = row[i]
        span = 1
        while i + span < len(row) and row[i + span] is None:
            span += 1
        cells.append((val or "", span))
        i += span
    return cells



def _cell_runs_xml(runs: list[TextRun], default_sz: int = 18) -> str:
    """Build w:r XML from runs, falling back to default_sz if no size info."""
    if not runs:
        return f'<w:r><w:rPr><w:sz w:val="{default_sz}"/><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/></w:rPr><w:t></w:t></w:r>'
    parts: list[str] = []
    for run in runs:
        if not run.text:
            continue
        safe = html.escape(run.text, quote=False)
        font = _word_font(run.font_name, run.bold)
        size = run.font_size_pt or 9.0
        sz = max(14, min(72, int(round(size * 2))))
        bold_xml = "<w:b/>" if run.bold or "Bold" in (run.font_name or "") else ""
        parts.append(
            f'<w:r><w:rPr>'
            f'<w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:cs="{font}"/>'
            f'<w:sz w:val="{sz}"/>'
            f'{bold_xml}'
            f'</w:rPr><w:t xml:space="preserve">{safe}</w:t></w:r>'
        )
    return "".join(parts) if parts else f'<w:r><w:rPr><w:sz w:val="{default_sz}"/></w:rPr><w:t></w:t></w:r>'


def _table_body_flow_xml(
    tblock: TableBlock,
    page_width_pt: float,
) -> str:
    """Body-level table with colspan support, proportional columns, and per-cell font runs."""
    rows = tblock.rows
    if not rows:
        return ""
    col_count = max(len(r) for r in rows)
    if col_count == 0:
        return ""

    def twip(pt: float) -> int:
        return max(1, int(pt * 20))

    # Compute column widths: use detected PDF widths if available, else equal split.
    col_widths_pt = tblock.col_widths_pt
    if col_widths_pt and len(col_widths_pt) == col_count and sum(col_widths_pt) > 0:
        total_w = sum(col_widths_pt)
        # Scale to usable page width (leave 36pt each side).
        usable = page_width_pt - 72
        scale = usable / total_w
        col_w_twips = [max(1, twip(w * scale)) for w in col_widths_pt]
    else:
        usable = page_width_pt - 72
        single = max(1, twip(usable) // col_count)
        col_w_twips = [single] * col_count

    tbl_w_twip = sum(col_w_twips)
    grid_xml = "".join(f'<w:gridCol w:w="{w}"/>' for w in col_w_twips)

    cell_runs_data = tblock.cell_runs  # list[list[list[TextRun]]]
    cell_aligns_data = tblock.cell_aligns  # list[list[str]]

    def cell_xml(text: str, colspan: int, ci_logical: int, runs: list[TextRun], align: str = "left") -> str:
        end = min(ci_logical + colspan, len(col_w_twips))
        cell_w = sum(col_w_twips[ci_logical:end]) if end > ci_logical else 1000
        grid_span = f'<w:gridSpan w:val="{colspan}"/>' if colspan > 1 else ""
        runs_xml = _cell_runs_xml(runs)
        jc = f'<w:jc w:val="center"/>' if align == "center" else ""
        return (
            f'<w:tc><w:tcPr>'
            f'<w:tcW w:w="{cell_w}" w:type="dxa"/>'
            f'{grid_span}'
            f'<w:tcBorders>'
            f'<w:top w:val="single" w:sz="4" w:color="000000"/>'
            f'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
            f'<w:left w:val="single" w:sz="4" w:color="000000"/>'
            f'<w:right w:val="single" w:sz="4" w:color="000000"/>'
            f'</w:tcBorders>'
            f'<w:shd w:val="clear" w:color="auto" w:fill="FFFFFF"/>'
            f'</w:tcPr>'
            f'<w:p><w:pPr><w:spacing w:before="0" w:after="0"/>{jc}</w:pPr>'
            f'{runs_xml}</w:p></w:tc>'
        )

    # First pass: decompose all rows into (text, span, runs, align) tuples.
    all_orig: list[list[tuple[str, int, list[TextRun], str]]] = []
    for ri, row in enumerate(rows):
        orig_row: list[tuple[str, int, list[TextRun], str]] = []
        ci = 0
        i = 0
        while i < len(row):
            span = 1
            while i + span < len(row) and row[i + span] is None:
                span += 1
            text = row[i] or ""
            runs: list[TextRun] = []
            align = "left"
            if cell_runs_data and ri < len(cell_runs_data) and ci < len(cell_runs_data[ri]):
                runs = cell_runs_data[ri][ci]
            if cell_aligns_data and ri < len(cell_aligns_data) and ci < len(cell_aligns_data[ri]):
                align = cell_aligns_data[ri][ci]
            orig_row.append((text, span, runs, align))
            ci += span
            i += span
        all_orig.append(orig_row)

    # Second pass: merge consecutive wide-span continuation rows.
    # Pattern: [('', 1, [], 'left'), (text, N>=4, runs, align)] where continuation has no bold first run.
    merged_orig: list[list[tuple[str, int, list[TextRun], str]]] = []
    for orig in all_orig:
        # Skip all-empty rows (PDF layout artifacts).
        if all(t == "" for t, _, _, _ in orig):
            continue
        is_wide = (len(orig) == 2 and orig[0][0] == "" and orig[0][1] == 1 and orig[1][1] >= 4)
        first_run_bold = bool(orig[1][2] and orig[1][2][0].bold) if is_wide else False
        is_continuation = is_wide and not first_run_bold
        if is_continuation and merged_orig:
            prev = merged_orig[-1]
            prev_is_wide = (len(prev) == 2 and prev[0][0] == "" and prev[0][1] == 1 and prev[1][1] >= 4)
            if prev_is_wide:
                prev_text, prev_span, prev_runs, prev_align = prev[1]
                curr_text, _, curr_runs, _ = orig[1]
                joined = (prev_text.rstrip() + " " + curr_text.strip()).strip()
                merged_orig[-1] = [prev[0], (joined, prev_span, prev_runs + curr_runs, prev_align)]
                continue
        merged_orig.append(list(orig))

    rows_xml_parts: list[str] = []
    for orig in merged_orig:
        # Merge leading empty x1 cell into adjacent wide content cell.
        if (len(orig) >= 2
                and orig[0][0] == ""
                and orig[0][1] == 1
                and orig[1][0].strip()
                and orig[1][1] >= 4):
            merged_span = orig[0][1] + orig[1][1]
            final = [(orig[1][0], merged_span, orig[1][2], orig[1][3])] + list(orig[2:])
        else:
            final = orig

        ci_logical = 0
        cells_xml = ""
        for text, colspan, runs, align in final:
            cells_xml += cell_xml(text, colspan, ci_logical, runs, align)
            ci_logical += colspan
        rows_xml_parts.append(f'<w:tr>{cells_xml}</w:tr>')

    rows_xml = "".join(rows_xml_parts)

    return (
        f'<w:tbl>'
        f'<w:tblPr>'
        f'<w:tblW w:w="{tbl_w_twip}" w:type="dxa"/>'
        f'<w:jc w:val="center"/>'
        f'<w:tblBorders>'
        f'<w:top w:val="single" w:sz="4" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
        f'<w:left w:val="single" w:sz="4" w:color="000000"/>'
        f'<w:right w:val="single" w:sz="4" w:color="000000"/>'
        f'<w:insideH w:val="single" w:sz="4" w:color="000000"/>'
        f'<w:insideV w:val="single" w:sz="4" w:color="000000"/>'
        f'</w:tblBorders>'
        f'</w:tblPr>'
        f'<w:tblGrid>{grid_xml}</w:tblGrid>'
        f'{rows_xml}'
        f'</w:tbl>'
    )


_NS_WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"


def _emit_text_para(doc: Document, block: TextBlock, page: Page) -> None:
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)

    # Determine alignment from bbox: if block center is past 40% of page width → center.
    block_center_x = block.bbox.x + block.bbox.w / 2
    if block_center_x > 0.40:
        pf.alignment = 1  # CENTER
    else:
        pf.alignment = 0  # LEFT

    for tr in block.runs:
        run = para.add_run(tr.text)
        run.bold = getattr(tr, "bold", False)
        font_size = max(8, int(tr.font_size_pt)) if getattr(tr, "font_size_pt", None) else 11
        run.font.size = Pt(font_size)
        run.font.name = _FALLBACK_FONT


def _emit_text_box(doc: Document, block: TextBlock, page: Page) -> None:
    """Emit a text block as a floating anchored text box positioned at exact PDF coordinates."""
    x_emu = int(block.bbox.x * page.width_pt * _PT_TO_EMU)
    y_emu = int(block.bbox.y * page.height_pt * _PT_TO_EMU)
    w_emu = max(int(block.bbox.w * page.width_pt * _PT_TO_EMU), 10000)
    h_emu = max(int(block.bbox.h * page.height_pt * _PT_TO_EMU), 10000)

    runs_xml = _runs_xml(block.runs)
    shape_id = _next_id()

    anchor_xml = (
        f'<w:drawing'
        f' xmlns:w="{_NS_W}"'
        f' xmlns:wp="{_NS_WP}"'
        f' xmlns:a="{_NS_A}"'
        f' xmlns:wps="{_NS_WPS}">'
        f'<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"'
        f' relativeHeight="251658240" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="page"><wp:posOffset>{x_emu}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="page"><wp:posOffset>{y_emu}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{w_emu}" cy="{h_emu}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:wrapNone/>'
        f'<wp:docPr id="{shape_id}" name="TextBox{shape_id}"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic>'
        f'<a:graphicData uri="{_NS_WPS}">'
        f'<wps:wsp>'
        f'<wps:cNvSpPr txBx="1"><a:spLocks noChangeArrowheads="1"/></wps:cNvSpPr>'
        f'<wps:spPr>'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{w_emu}" cy="{h_emu}"/></a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'<a:noFill/>'
        f'<a:ln><a:noFill/></a:ln>'
        f'</wps:spPr>'
        f'<wps:txbx>'
        f'<w:txbxContent>'
        f'<w:p><w:pPr><w:spacing w:before="0" w:after="0"/></w:pPr>'
        f'{runs_xml}'
        f'</w:p>'
        f'</w:txbxContent>'
        f'</wps:txbx>'
        f'<wps:bodyPr rot="0" spcFirstLastPara="0" vertOverflow="clip" horzOverflow="clip"'
        f' vert="horz" wrap="none" lIns="0" tIns="0" rIns="0" bIns="0" anchor="t" anchorCtr="0">'
        f'<a:normAutofit/>'
        f'</wps:bodyPr>'
        f'</wps:wsp>'
        f'</a:graphicData>'
        f'</a:graphic>'
        f'</wp:anchor>'
        f'</w:drawing>'
    )

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run()
    run._r.append(parse_xml(anchor_xml))


def _add_background_image(doc: Document, bg_path: Path, page: Page) -> None:
    """Insert full-page masked background as a behindDoc=1 anchor in a zero-height paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run()
    run.add_picture(str(bg_path), width=Pt(page.width_pt), height=Pt(page.height_pt))

    # Extract the relationship id from the inline picture, then replace the entire
    # <w:drawing> with a new one that uses a behindDoc anchor instead of inline.
    r_elem = run._r
    drawing = r_elem.find(qn("w:drawing"))
    if drawing is None:
        return
    inline = drawing.find(qn("wp:inline"))
    if inline is None:
        return

    ns_a = "http://schemas.openxmlformats.org/drawingml/2006/main"
    ns_r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    blip = inline.find(f".//{{{ns_a}}}blip")
    if blip is None:
        return
    rel_id = blip.get(f"{{{ns_r}}}embed")
    if not rel_id:
        return

    # Replace the whole <w:drawing> element with a new anchor-based one.
    shape_id = _next_id()
    anchor_xml = _anchor_image_xml(rel_id, page.width_pt, page.height_pt, shape_id)
    new_drawing = parse_xml(anchor_xml)
    r_elem.remove(drawing)
    r_elem.append(new_drawing)


def emit_docx_native_fidelity(
    ir: DocumentIR,
    output_path: Path,
    *,
    page_backgrounds: list[Path],
    mask_dir: Path | None = None,
    dpi: int = 200,
) -> Path:
    """Masked page raster (borders, logo) + editable text table with PDF fonts."""
    global _SHAPE_ID
    _SHAPE_ID = 2000

    doc = Document()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    work_mask = mask_dir or output_path.parent / "masked_assets"
    work_mask.mkdir(parents=True, exist_ok=True)

    pages = ir.pages or [Page(index=i) for i in range(ir.source.page_count)]
    page_count = ir.source.page_count

    _W_NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'

    for pi in range(page_count):
        page = pages[pi] if pi < len(pages) else Page(index=pi)
        if pi > 0:
            doc.add_section()
        _set_section(doc.sections[-1], page)

        page_texts = _page_text_blocks(ir, pi)
        page_tables = _page_table_blocks(ir, pi)

        table_top_y = min((t.bbox.y for t in page_tables), default=1.0)
        table_bot_y = max((t.bbox.y + t.bbox.h for t in page_tables), default=0.0)

        # Insert background image if provided for this page.
        bg_path = page_backgrounds[pi] if pi < len(page_backgrounds) else None
        if bg_path and bg_path.exists():
            _add_background_image(doc, bg_path, page)

        if page_tables:
            # Document-style page: emit text as in-flow paragraphs, tables as Word tables.
            for block in sorted(page_texts, key=lambda b: b.bbox.y):
                if block.bbox.y >= table_top_y:
                    break
                _emit_text_para(doc, block, page)

            for tblock in page_tables:
                if not tblock.rows:
                    continue
                inner = _table_body_flow_xml(tblock, page.width_pt)
                if inner:
                    wrapped = inner.replace("<w:tbl>", f"<w:tbl {_W_NS}>", 1)
                    doc.element.body.append(parse_xml(wrapped))

            for block in sorted(page_texts, key=lambda b: b.bbox.y):
                if block.bbox.y < table_bot_y:
                    continue
                _emit_text_para(doc, block, page)
        else:
            # Slide-style page (no tables): emit all text as positioned floating text boxes.
            for block in sorted(page_texts, key=lambda b: b.bbox.y):
                _emit_text_box(doc, block, page)

    doc.save(output_path)
    return output_path
