"""DOCX editable layer — OCR text + figures per PDF page section."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt

from src.pipeline.cleanup import clean_ocr_text
from src.pipeline.gap_fill import _is_gap_path
from src.pipeline.ir import DocumentIR, ImageBlock, PageBreakBlock, TableBlock, TextBlock

_GARBAGE = re.compile(r"^[\s\":;.,0-9«»\-–—]{0,4}$")


def _is_garbage(text: str) -> bool:
    t = text.strip()
    if len(t) < 2:
        return True
    return bool(_GARBAGE.match(t))


def _sort_key(block) -> tuple[float, float]:
    bb = block.bbox
    return (bb.y, bb.x)


def _page_blocks(ir: DocumentIR, page_index: int) -> list:
    blocks = []
    for block in ir.blocks:
        if isinstance(block, PageBreakBlock):
            continue
        if isinstance(block, ImageBlock):
            if _is_gap_path(block.image_path):
                continue
            if not block.image_path or not Path(block.image_path).exists():
                continue
        if getattr(block, "page_index", -1) == page_index:
            blocks.append(block)
    blocks.sort(key=_sort_key)
    return blocks


def _flush_text_buf(doc: Document, buf: list[str]) -> None:
    if not buf:
        return
    doc.add_paragraph(" ".join(buf))
    buf.clear()


def _add_image(doc: Document, block: ImageBlock) -> None:
    path = Path(block.image_path)
    if not path.exists():
        return
    width = Inches(min(6.0, max(1.5, block.bbox.w * 6.5)))
    try:
        doc.add_paragraph().add_run().add_picture(str(path), width=width)
    except Exception:
        doc.add_paragraph(f"[Изображение: {path.name}]")


def emit_docx_editable(ir: DocumentIR, output_path: Path) -> Path:
    """32 sections «Страница N»: текст, таблицы, иллюстрации (печати, подписи, рисунки)."""
    doc = Document()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc.styles["Normal"].font.size = Pt(11)
    page_count = ir.source.page_count

    doc.add_paragraph(
        f"Распознанный текст и иллюстрации по {page_count} страницам исходного PDF. "
        "Вёрстка оригинала не сохраняется — только содержимое по страницам."
    )

    for pi in range(page_count):
        if pi > 0:
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.PAGE)

        doc.add_heading(f"Страница {pi + 1}", level=1)

        blocks = _page_blocks(ir, pi)
        if not blocks:
            doc.add_paragraph("(нет распознанного содержимого)")
            continue

        text_buf: list[str] = []
        last_y = -1.0

        for block in blocks:
            if isinstance(block, TextBlock):
                text = clean_ocr_text(block.text)
                if _is_garbage(text):
                    continue
                y = block.bbox.y
                if text_buf and abs(y - last_y) > 0.04:
                    _flush_text_buf(doc, text_buf)
                text_buf.append(text)
                last_y = y

            elif isinstance(block, TableBlock):
                _flush_text_buf(doc, text_buf)
                rows = block.rows
                if rows:
                    cols = max(len(r) for r in rows)
                    table = doc.add_table(rows=len(rows), cols=cols)
                    table.style = "Table Grid"
                    for ri, row in enumerate(rows):
                        for ci in range(cols):
                            cell = row[ci] if ci < len(row) else ""
                            table.rows[ri].cells[ci].text = clean_ocr_text(cell)
                    doc.add_paragraph()

            elif isinstance(block, ImageBlock):
                _flush_text_buf(doc, text_buf)
                if block.caption:
                    doc.add_paragraph(block.caption, style="Caption")
                _add_image(doc, block)

        _flush_text_buf(doc, text_buf)

    doc.save(output_path)
    return output_path
