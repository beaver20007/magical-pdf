"""DOCX visual layer — exactly N pages, each page is one full-scan image."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

from src.pipeline.ir import DocumentIR, Page


def emit_docx_visual(
    ir: DocumentIR,
    output_path: Path,
    *,
    page_backgrounds: list[Path],
) -> Path:
    """One page break + one full-page image per PDF page. No textboxes."""
    doc = Document()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    pages = ir.pages or [Page(index=i) for i in range(ir.source.page_count)]

    for pi, page in enumerate(pages):
        if pi > 0:
            brk = doc.add_paragraph()
            brk.paragraph_format.space_before = Pt(0)
            brk.paragraph_format.space_after = Pt(0)
            brk.add_run().add_break(WD_BREAK.PAGE)

        if pi == 0:
            section = doc.sections[0]
        else:
            section = doc.sections[-1]

        section.page_width = Pt(page.width_pt)
        section.page_height = Pt(page.height_pt)
        section.top_margin = Pt(0)
        section.bottom_margin = Pt(0)
        section.left_margin = Pt(0)
        section.right_margin = Pt(0)

        bg = page_backgrounds[pi] if pi < len(page_backgrounds) else None
        if not bg or not bg.exists():
            doc.add_paragraph(f"[Missing page image {pi + 1}]")
            continue

        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = Pt(page.height_pt)
        run = para.add_run()
        run.add_picture(str(bg), width=Pt(page.width_pt), height=Pt(page.height_pt))

    doc.save(output_path)
    return output_path
