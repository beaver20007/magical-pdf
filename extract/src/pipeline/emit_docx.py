"""DocumentIR → editable DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt

from src.pipeline.cleanup import clean_ocr_text
from src.pipeline.ir import (
    DocumentIR,
    ImageBlock,
    PageBreakBlock,
    TableBlock,
    TextBlock,
)

_ROLE_STYLE = {
    "title": "Title",
    "heading": "Heading 1",
    "paragraph": "Normal",
    "list_item": "List Bullet",
    "caption": "Caption",
    "footer": "Footer",
    "header": "Header",
    "unknown": "Normal",
}


def _insert_page_break(doc: Document) -> None:
    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)


def emit_docx(ir: DocumentIR, output_path: Path) -> Path:
    doc = Document()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    current_page = -1

    for block in ir.blocks:
        if isinstance(block, PageBreakBlock):
            _insert_page_break(doc)
            continue

        page_idx = getattr(block, "page_index", 0)
        if page_idx > current_page and current_page >= 0:
            _insert_page_break(doc)
        current_page = max(current_page, page_idx)

        if isinstance(block, TextBlock):
            text = clean_ocr_text(block.text)
            style = _ROLE_STYLE.get(block.role, "Normal")
            try:
                para = doc.add_paragraph(text, style=style)
            except KeyError:
                para = doc.add_paragraph(block.text)
            if block.role in ("title", "heading"):
                for run in para.runs:
                    run.bold = True
            continue

        if isinstance(block, TableBlock):
            rows = block.rows
            if not rows:
                continue
            col_count = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=col_count)
            table.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci in range(col_count):
                    cell_text = row[ci] if ci < len(row) else ""
                    table.rows[ri].cells[ci].text = clean_ocr_text(cell_text)
            doc.add_paragraph()
            continue

        if isinstance(block, ImageBlock):
            if "gap_fills" in (block.image_path or "").replace("\\", "/"):
                continue
            if block.caption:
                doc.add_paragraph(block.caption, style="Caption")
            if block.image_path and Path(block.image_path).exists():
                width = Inches(5.0)
                if block.bbox.w > 0:
                    width = Inches(min(6.5, max(2.0, block.bbox.w * 6.5)))
                try:
                    doc.add_picture(block.image_path, width=width)
                except Exception:
                    doc.add_paragraph(f"[Image: {block.image_path}]")
            doc.add_paragraph()

    if not doc.paragraphs and not doc.tables:
        doc.add_paragraph("(No content extracted)")

    doc.save(output_path)
    return output_path
