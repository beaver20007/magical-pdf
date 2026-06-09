"""DOCX: N pages = N PDF pages. Each page: scan image + optional text overlays in ONE paragraph."""

from __future__ import annotations

import html
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.shared import Pt

from src.pipeline.cleanup import clean_ocr_text
from src.pipeline.font_estimate import DEFAULT_FONT_NAME, estimate_font_size_pt
from src.pipeline.ir import (
    DocumentIR,
    ImageBlock,
    Page,
    TableBlock,
    TextBlock,
)


def _set_section_size(section, page: Page) -> None:
    section.page_width = Pt(page.width_pt)
    section.page_height = Pt(page.height_pt)
    section.top_margin = Pt(0)
    section.bottom_margin = Pt(0)
    section.left_margin = Pt(0)
    section.right_margin = Pt(0)


def _bbox_to_pt(bbox, page: Page) -> tuple[float, float, float, float]:
    left = bbox.x * page.width_pt
    top = bbox.y * page.height_pt
    width = max(12.0, bbox.w * page.width_pt)
    height = max(8.0, bbox.h * page.height_pt)
    return left, top, width, height


def _vml_textbox_xml(
    text: str,
    left_pt: float,
    top_pt: float,
    width_pt: float,
    height_pt: float,
    font_size_pt: float,
    font_name: str,
) -> str:
    safe = html.escape(text, quote=False)
    sz_half = max(16, min(28, int(round(font_size_pt * 2))))
    return f"""<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office">
  <v:shape type="#_x0000_t202" stroked="f" filled="t" fillcolor="white" o:allowincell="f"
    style="position:absolute;margin-left:{left_pt:.2f}pt;margin-top:{top_pt:.2f}pt;width:{width_pt:.2f}pt;height:{height_pt:.2f}pt;z-index:251658240;mso-position-horizontal-relative:page;mso-position-vertical-relative:page">
    <v:textbox inset="1,1,1,1">
      <w:txbxContent>
        <w:p>
          <w:pPr><w:spacing w:before="0" w:after="0"/></w:pPr>
          <w:r>
            <w:rPr>
              <w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>
              <w:sz w:val="{sz_half}"/>
            </w:rPr>
            <w:t xml:space="preserve">{safe}</w:t>
          </w:r>
        </w:p>
      </w:txbxContent>
    </v:textbox>
  </v:shape>
</w:pict>"""


def _text_blocks_on_page(ir: DocumentIR, page_index: int) -> list[TextBlock | TableBlock]:
    out: list[TextBlock | TableBlock] = []
    for b in ir.blocks:
        if getattr(b, "page_index", -1) != page_index or b.type == "page_break":
            continue
        if isinstance(b, (TextBlock, TableBlock)):
            out.append(b)
    return out


def emit_docx_scans_only(
    ir: DocumentIR,
    output_path: Path,
    *,
    page_backgrounds: list[Path],
) -> Path:
    """32 pages = 32 scans only. Guaranteed visual match, text not overlaid."""
    doc = Document()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    pages = ir.pages or [Page(index=i) for i in range(ir.source.page_count)]

    for pi, page in enumerate(pages):
        if pi > 0:
            doc.add_section()
        _set_section_size(doc.sections[-1], page)
        bg = page_backgrounds[pi] if pi < len(page_backgrounds) else None
        if bg and bg.exists():
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.add_run().add_picture(str(bg), width=Pt(page.width_pt), height=Pt(page.height_pt))

    doc.save(output_path)
    return output_path


def emit_docx_positioned(
    ir: DocumentIR,
    output_path: Path,
    *,
    page_backgrounds: list[Path] | None = None,
    overlay_text: bool = True,
) -> Path:
    """
    One paragraph per PDF page (prevents 64-page doubling).

    overlay_text=False → scan-only (recommended if Word mangles VML).
    """
    if not overlay_text or not page_backgrounds:
        if not page_backgrounds:
            raise ValueError("page_backgrounds required")
        return emit_docx_scans_only(ir, output_path, page_backgrounds=page_backgrounds)

    doc = Document()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    pages = ir.pages or [Page(index=i) for i in range(ir.source.page_count)]

    for pi, page in enumerate(pages):
        if pi > 0:
            doc.add_section()
        _set_section_size(doc.sections[-1], page)

        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(0.1)

        run = para.add_run()

        bg = page_backgrounds[pi] if pi < len(page_backgrounds) else None
        if bg and bg.exists():
            run.add_picture(str(bg), width=Pt(page.width_pt), height=Pt(page.height_pt))

        for block in _text_blocks_on_page(ir, page.index):
            left, top, width, height = _bbox_to_pt(block.bbox, page)

            if isinstance(block, TextBlock):
                text = clean_ocr_text(block.text)
                if not text.strip():
                    continue
                font_pt = block.font_size_pt or estimate_font_size_pt(
                    block.bbox, page, text
                )
                font_name = block.font_name or DEFAULT_FONT_NAME
                run._r.append(
                    parse_xml(
                        _vml_textbox_xml(
                            text, left, top, width, height, font_pt, font_name
                        )
                    )
                )

            elif isinstance(block, TableBlock):
                rows = block.rows
                if not rows:
                    continue
                cols = max(len(r) for r in rows)
                table_text = "\n".join(
                    " | ".join(row[ci] if ci < len(row) else "" for ci in range(cols))
                    for row in rows
                )
                table_text = clean_ocr_text(table_text)
                font_pt = estimate_font_size_pt(block.bbox, page, table_text) * 0.9
                run._r.append(
                    parse_xml(
                        _vml_textbox_xml(
                            table_text,
                            left,
                            top,
                            width,
                            height,
                            font_pt,
                            DEFAULT_FONT_NAME,
                        )
                    )
                )

        # Skip ImageBlock — scan background already contains stamps/signatures

    doc.save(output_path)
    return output_path
